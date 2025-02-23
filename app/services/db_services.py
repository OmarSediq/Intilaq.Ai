from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from sqlalchemy.sql import delete
from sqlalchemy.sql import text
from app.models import User, ResetCode,Header,Education,Experience,Objective,VolunteeringExperience,SkillsLanguages,Awards,Certifications,Projects,LoginAttempt
from app.dependencies import get_password_hash
from datetime import datetime, timedelta, timezone
from fastapi import HTTPException
import json
import io
from docx import Document
from docx.shared import Pt
from bs4 import BeautifulSoup
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from app.utils.jwt_utils import decode_access_token
from fastapi import Depends, HTTPException

async def create_user(username: str, password: str, email: str, db: AsyncSession):
    """
    Creates a new user and ensures no duplicate username or email exists.
    """
    result = await db.execute(
        select(User).filter((User.username == username) | (User.email == email))
    )
    existing_user = result.scalars().first()

    if existing_user:
        raise HTTPException(status_code=400, detail="Username or email already exists")

    hashed_password = get_password_hash(password)

    user = User(username=username, email=email, hashed_password=hashed_password, is_verified=0)
    db.add(user)
    await db.commit()
    await db.refresh(user)
    return user

async def get_user_by_username(username: str, db: AsyncSession):
    """
    Retrieves a user by their username.
    """
    result = await db.execute(select(User).filter(User.username == username))
    user = result.scalars().first()
    return user

async def get_user_by_email(email: str, db: AsyncSession):
    """
    Retrieves a user by their email address.
    """
    result = await db.execute(select(User).filter(User.email == email))
    user = result.scalars().first()
    return user

async def save_reset_code(email: str, code: str, db: AsyncSession):
    """
    Saves a reset code for the given email, ensuring no duplicate entries.
    """
    expiration_time = datetime.now(timezone.utc) + timedelta(minutes=5)
    expiration_time = expiration_time.replace(tzinfo=None)  # Remove timezone info for compatibility

    reset_code = ResetCode(email=email, code=code, created_at=expiration_time)

    # Delete existing reset codes for the email
    await db.execute(delete(ResetCode).where(ResetCode.email == email))

    db.add(reset_code)
    try:
        await db.commit()
    except Exception as e:
        await db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

async def verify_reset_code(email: str, code: str, db: AsyncSession) -> bool:
    """
    Verifies the reset code for the given email.
    """
    result = await db.execute(select(ResetCode).filter(ResetCode.email == email, ResetCode.code == code))
    reset_code = result.scalars().first()

    if reset_code and reset_code.created_at > datetime.now(timezone.utc):
        return True

    return False

async def update_verification_status(email: str, db: AsyncSession):
    """
    Updates the verification status of the user associated with the given email.
    """
    user = await get_user_by_email(email, db)
    if user:
        user.is_verified = 1
        try:
            await db.commit()
            await db.refresh(user)
        except Exception as e:
            await db.rollback()
            raise HTTPException(status_code=500, detail=f"Failed to update verification status: {str(e)}")
        return user
    raise HTTPException(status_code=404, detail="User not found")

async def update_user_details(user_id: int, updated_data: dict, db: AsyncSession):
    """
    Updates user details by user ID.
    """
    result = await db.execute(select(User).filter(User.id == user_id))
    user = result.scalars().first()

    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    for key, value in updated_data.items():
        setattr(user, key, value)

    try:
        await db.commit()
        await db.refresh(user)
    except Exception as e:
        await db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to update user details: {str(e)}")

    return user

async def delete_user_by_id(user_id: int, db: AsyncSession):
    """
    Deletes a user by their ID.
    """
    result = await db.execute(select(User).filter(User.id == user_id))
    user = result.scalars().first()

    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    await db.delete(user)
    try:
        await db.commit()
    except Exception as e:
        await db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to delete user: {str(e)}")

    return True

async def get_user_by_id(user_id: int, db: AsyncSession):
    result = await db.execute(select(User).filter(User.id == user_id, User.is_active == 1))
    user = result.scalars().first()

    if not user:
        raise HTTPException(status_code=404, detail="User not found or inactive")

    return user

async def generate_docx_from_html(html_content):
   
    soup = BeautifulSoup(html_content, "html.parser")

    doc = Document()

    def add_paragraph(text, bold=False, font_size=12, alignment=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)
        if alignment:
            p.alignment = alignment

    for tag in soup.find_all(["h1", "h2", "h3", "p", "ul", "li", "a"]):
        if tag.name == "h1":
            add_paragraph(tag.get_text(), bold=True, font_size=16, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
        elif tag.name == "h2":
            add_paragraph(tag.get_text(), bold=True, font_size=14)
        elif tag.name == "h3":
            add_paragraph(tag.get_text(), bold=True, font_size=12)
        elif tag.name == "p":
            add_paragraph(tag.get_text(), font_size=11)
        elif tag.name == "ul":
            for li in tag.find_all("li"):
                p = doc.add_paragraph("• " + li.get_text(), style="List Bullet")
                p.paragraph_format.space_after = Pt(6)
        elif tag.name == "a":
            add_paragraph(f"{tag.get_text()} ({tag['href']})", font_size=11)

    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer


async def get_user_by_header_id(db: AsyncSession, header_id: int):
    query = text("""
       SELECT 
    h.full_name AS full_name, h.job_title, h.email, h.phone_number AS phone_number, h.address,
    h.years_of_experience, h.linkedin_profile, h.github_profile,
    COALESCE(o.description, '') AS objective, 

    COALESCE(json_agg(DISTINCT e.*) FILTER (WHERE e.id IS NOT NULL), '[]') AS education,
    COALESCE(json_agg(DISTINCT exp.*) FILTER (WHERE exp.id IS NOT NULL), '[]') AS experience,
    COALESCE(json_agg(DISTINCT p.*) FILTER (WHERE p.id IS NOT NULL), '[]') AS projects,
    COALESCE(json_agg(DISTINCT v.*) FILTER (WHERE v.id IS NOT NULL), '[]') AS volunteering_experience,

   COALESCE(json_agg(DISTINCT jsonb_build_object(
    'skill', s.skills
)) FILTER (WHERE s.skills IS NOT NULL AND s.skills <> ''), '[]') AS technical_skills,

COALESCE(json_agg(DISTINCT jsonb_build_object(
    'language', s.languages,
    'level', s.level
)) FILTER (WHERE s.languages IS NOT NULL AND s.languages <> ''), '[]') AS languages,

    COALESCE(json_agg(DISTINCT jsonb_build_object(
        'certification_title', c.certification_title, 
        'link', c.link
    )) FILTER (WHERE c.id IS NOT NULL), '[]') AS certifications,

    COALESCE(json_agg(DISTINCT jsonb_build_object(
        'award', a.award, 
        'organization', a.organization, 
        'start_date', a.start_date, 
        'end_date', a.end_date
    )) FILTER (WHERE a.id IS NOT NULL), '[]') AS awards

FROM header h
LEFT JOIN objective o ON h.id = o.header_id
LEFT JOIN education e ON h.id = e.header_id
LEFT JOIN experience exp ON h.id = exp.header_id
LEFT JOIN projects p ON h.id = p.header_id
LEFT JOIN skills_languages s ON h.id = s.header_id  
LEFT JOIN volunteering_experience v ON h.id = v.header_id
LEFT JOIN certifications c ON h.id = c.header_id
LEFT JOIN awards a ON h.id = a.header_id
WHERE h.id = :header_id
GROUP BY h.id, o.description;
    """)

    print(f"Executing query for header_id: {header_id}")  

    result = await db.execute(query, {"header_id": header_id})
    user = result.mappings().first()  

    print("Raw Data from Database:", user)  

    if not user:
        print("No user found!")  
        return None  

    user_dict = dict(user)

    print("User Dictionary Before Processing JSON:", user_dict)  

    for key in ["education", "experience", "projects", "volunteering_experience", "certifications", "awards"]:
      if isinstance(user_dict.get(key), str):
        user_dict[key] = json.loads(user_dict[key])

        if isinstance(user_dict.get("technical_skills"), str):
                user_dict["technical_skills"] = json.loads(user_dict["technical_skills"])

        if isinstance(user_dict.get("languages"), str):
                user_dict["languages"] = json.loads(user_dict["languages"])


    print("User Dictionary After Processing JSON:", user_dict)  

    return user_dict

# --------------------- Header Management --------------------- #

async def create_header(data: dict, db: AsyncSession):
    header = Header(**data)
    db.add(header)
    await db.commit()
    await db.refresh(header)
    return header

async def get_header_by_id(header_id: int, db: AsyncSession):
    result = await db.execute(select(Header).filter(Header.id == header_id))
    return result.scalars().first()

async def update_header(header_id: int, data: dict, db: AsyncSession):
    header = await get_header_by_id(header_id, db)
    if not header:
        raise HTTPException(status_code=404, detail="Header not found")
    for key, value in data.items():
        setattr(header, key, value)
    await db.commit()
    await db.refresh(header)
    return header

async def delete_header(header_id: int, db: AsyncSession):
    header = await get_header_by_id(header_id, db)
    if not header:
        raise HTTPException(status_code=404, detail="Header not found")
    await db.delete(header)
    await db.commit()
    return True

# --------------------- Education Management --------------------- #

async def create_education(data: dict, db: AsyncSession):
    education = Education(**data)
    db.add(education)
    await db.commit()
    await db.refresh(education)
    return education

async def get_education_by_id(education_id: int, db: AsyncSession):
    result = await db.execute(select(Education).filter(Education.id == education_id))
    return result.scalars().first()

async def update_education(education_id: int, data: dict, db: AsyncSession):
    education = await get_education_by_id(education_id, db)
    if not education:
        raise HTTPException(status_code=404, detail="Education not found")
    for key, value in data.items():
        setattr(education, key, value)
    await db.commit()
    await db.refresh(education)
    return education

async def delete_education(education_id: int, db: AsyncSession):
    education = await get_education_by_id(education_id, db)
    if not education:
        raise HTTPException(status_code=404, detail="Education not found")
    await db.delete(education)
    await db.commit()
    return True

# --------------------- Experience Management --------------------- #

async def create_experience(data: dict, db: AsyncSession):
    experience = Experience(**data)
    db.add(experience)
    await db.commit()
    await db.refresh(experience)
    return experience

async def get_experience_by_id(experience_id: int, db: AsyncSession):
    result = await db.execute(select(Experience).filter(Experience.id == experience_id))
    return result.scalars().first()

async def update_experience(experience_id: int, data: dict, db: AsyncSession):
    experience = await get_experience_by_id(experience_id, db)
    if not experience:
        raise HTTPException(status_code=404, detail="Experience not found")
    for key, value in data.items():
        setattr(experience, key, value)
    await db.commit()
    await db.refresh(experience)
    return experience

async def delete_experience(experience_id: int, db: AsyncSession):
    experience = await get_experience_by_id(experience_id, db)
    if not experience:
        raise HTTPException(status_code=404, detail="Experience not found")
    await db.delete(experience)
    await db.commit()
    return True

# --------------------- Volunteering Management --------------------- #

async def create_volunteering(data: dict, db: AsyncSession):
    volunteering = VolunteeringExperience(**data)
    db.add(volunteering)
    await db.commit()
    await db.refresh(volunteering)
    return volunteering

async def get_volunteering_by_id(volunteering_id: int, db: AsyncSession):
    result = await db.execute(select(VolunteeringExperience).filter(VolunteeringExperience.id == volunteering_id))
    return result.scalars().first()

async def update_volunteering(volunteering_id: int, data: dict, db: AsyncSession):
    volunteering = await get_volunteering_by_id(volunteering_id, db)
    if not volunteering:
        raise HTTPException(status_code=404, detail="Volunteering experience not found")
    for key, value in data.items():
        setattr(volunteering, key, value)
    await db.commit()
    await db.refresh(volunteering)
    return volunteering

async def delete_volunteering(volunteering_id: int, db: AsyncSession):
    volunteering = await get_volunteering_by_id(volunteering_id, db)
    if not volunteering:
        raise HTTPException(status_code=404, detail="Volunteering experience not found")
    await db.delete(volunteering)
    await db.commit()
    return True


# --------------------- Skills & Languages Management --------------------- #

async def create_skills_languages(data: dict, db: AsyncSession):
    skills_languages = SkillsLanguages(**data)
    db.add(skills_languages)
    await db.commit()
    await db.refresh(skills_languages)
    return skills_languages

async def get_skills_languages_by_id(skills_languages_id: int, db: AsyncSession):
    result = await db.execute(select(SkillsLanguages).filter(SkillsLanguages.id == skills_languages_id))
    return result.scalars().first()

async def update_skills_languages(skills_languages_id: int, data: dict, db: AsyncSession):
    skills_languages = await get_skills_languages_by_id(skills_languages_id, db)
    if not skills_languages:
        raise HTTPException(status_code=404, detail="Skills & Languages not found")
    for key, value in data.items():
        setattr(skills_languages, key, value)
    await db.commit()
    await db.refresh(skills_languages)
    return skills_languages

async def delete_skills_languages(skills_languages_id: int, db: AsyncSession):
    skills_languages = await get_skills_languages_by_id(skills_languages_id, db)
    if not skills_languages:
        raise HTTPException(status_code=404, detail="Skills & Languages not found")
    await db.delete(skills_languages)
    await db.commit()
    return True


# --------------------- Projects Management --------------------- #
#
async def create_project(data: dict, db: AsyncSession):
    project = Projects(**data)
    db.add(project)
    await db.commit()
    await db.refresh(project)
    return project

async def get_project_by_id(project_id: int, db: AsyncSession):
    result = await db.execute(select(Projects).filter(Projects.id == project_id))
    return result.scalars().first()

async def update_project(project_id: int, data: dict, db: AsyncSession):
    project = await get_project_by_id(project_id, db)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    for key, value in data.items():
        setattr(project, key, value)
    await db.commit()
    await db.refresh(project)
    return project

async def delete_project(project_id: int, db: AsyncSession):
    project = await get_project_by_id(project_id, db)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    await db.delete(project)
    await db.commit()
    return True

# --------------------- Certifications Management --------------------- #

async def create_certification(data: dict, db: AsyncSession):
    certification = Certifications(**data)
    db.add(certification)
    await db.commit()
    await db.refresh(certification)
    return certification

async def get_certification_by_id(certification_id: int, db: AsyncSession):
    result = await db.execute(select(Certifications).filter(Certifications.id == certification_id))
    return result.scalars().first()

async def update_certification(certification_id: int, data: dict, db: AsyncSession):
    certification = await get_certification_by_id(certification_id, db)
    if not certification:
        raise HTTPException(status_code=404, detail="Certification not found")
    for key, value in data.items():
        setattr(certification, key, value)
    await db.commit()
    await db.refresh(certification)
    return certification

async def delete_certification(certification_id: int, db: AsyncSession):
    certification = await get_certification_by_id(certification_id, db)
    if not certification:
        raise HTTPException(status_code=404, detail="Certification not found")
    await db.delete(certification)
    await db.commit()
    return True

# --------------------- Awards Management --------------------- #

async def create_award(data: dict, db: AsyncSession):
    award = Awards(**data)
    db.add(award)
    await db.commit()
    await db.refresh(award)
    return award

async def get_award_by_id(award_id: int, db: AsyncSession):
    result = await db.execute(select(Awards).filter(Awards.id == award_id))
    return result.scalars().first()

async def update_award(award_id: int, data: dict, db: AsyncSession):
    award = await get_award_by_id(award_id, db)
    if not award:
        raise HTTPException(status_code=404, detail="Award not found")
    for key, value in data.items():
        setattr(award, key, value)
    await db.commit()
    await db.refresh(award)
    return award

async def delete_award(award_id: int, db: AsyncSession):
    award = await get_award_by_id(award_id, db)
    if not award:
        raise HTTPException(status_code=404, detail="Award not found")
    await db.delete(award)
    await db.commit()
    return True

# --------------------- Objective Management --------------------- #

async def create_objective(data: dict, db: AsyncSession):
    objective = Objective(**data)
    db.add(objective)
    await db.commit()
    await db.refresh(objective)
    return objective

async def get_objective_by_id(objective_id: int, db: AsyncSession):
    result = await db.execute(select(Objective).filter(Objective.id == objective_id))
    return result.scalars().first()

async def update_objective(objective_id: int, data: dict, db: AsyncSession):
    objective = await get_objective_by_id(objective_id, db)
    if not objective:
        raise HTTPException(status_code=404, detail="Objective not found")
    for key, value in data.items():
        setattr(objective, key, value)
    await db.commit()
    await db.refresh(objective)
    return objective

async def delete_objective(objective_id: int, db: AsyncSession):
    objective = await get_objective_by_id(objective_id, db)
    if not objective:
        raise HTTPException(status_code=404, detail="Objective not found")
    await db.delete(objective)
    await db.commit()
    return True