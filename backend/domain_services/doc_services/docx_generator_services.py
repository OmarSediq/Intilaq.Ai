from docx import Document
from bs4 import BeautifulSoup
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io

class DocxGenerator:
    @staticmethod
    async def generate_from_html(html_content: str):
        soup = BeautifulSoup(html_content, "html.parser")
        doc = Document()

        def add_paragraph(text, bold=False, font_size=12, alignment=None):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(font_size)
            if alignment:
                p.alignment = alignment

        for tag in soup.find_all(["h1", "h2", "h3", "p", "ul", "li", "a"]):
            if tag.name == "h1":
                add_paragraph(tag.get_text(), bold=True, font_size=16, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
            elif tag.name == "h2":
                add_paragraph(tag.get_text(), bold=True, font_size=14)
            elif tag.name == "h3":
                add_paragraph(tag.get_text(), bold=True, font_size=12)
            elif tag.name == "p":
                add_paragraph(tag.get_text(), font_size=11)
            elif tag.name == "ul":
                for li in tag.find_all("li"):
                    p = doc.add_paragraph("• " + li.get_text(), style="List Bullet")
                    p.paragraph_format.space_after = Pt(6)
            elif tag.name == "a":
                add_paragraph(f"{tag.get_text()} ({tag['href']})", font_size=11)

        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        return docx_buffer
